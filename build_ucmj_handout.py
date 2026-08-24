#!/usr/bin/env python3
"""Build print-ready PDF and editable DOCX handouts from UCMJ_Study_Guide.md.

One parsed model feeds both renderers so the two deliverables cannot drift.
"""
import re, os, sys, html as _html

SRC   = os.path.join(os.path.dirname(os.path.abspath(__file__)), "UCMJ_Study_Guide.md")
OUT   = os.path.dirname(os.path.abspath(__file__))
EXAM  = "071-OAXXD012"
TITLE = "Military Justice / UCMJ Study Guide"

# Tier descriptions, verbatim from the source preamble bullets, reused as
# subtitles so each tier page states its own purpose without a flip back to p1.
TIER_SUB = {
 "TIER 1": "Tier 1 is the hard-recall material. Numbered lists, fixed categories, "
           "specific counts. This is what multiple choice exams are built out of. "
           "Know all of it cold.",
 "TIER 2": 'Tier 2 is concept material. Likely tested, but usually as "which of '
           'these is an example of" rather than pure recall.',
 "TIER 3": "Tier 3 is background and reference. Lower probability, but cheap to skim "
           "and it makes Tier 1 make sense.",
}

# ---------------------------------------------------------------- inline parse
def parse_inline(s):
    """-> [(text, bold, italic)] handling **bold** and *italic*."""
    out, pos = [], 0
    for m in re.finditer(r'\*\*(.+?)\*\*|\*(.+?)\*', s):
        if m.start() > pos:
            out.append((s[pos:m.start()], False, False))
        if m.group(1) is not None:
            out.append((m.group(1), True, False))
        else:
            out.append((m.group(2), False, True))
        pos = m.end()
    if pos < len(s):
        out.append((s[pos:], False, False))
    return [r for r in out if r[0]]

# ----------------------------------------------------------------- block parse
def parse(md):
    lines = md.split("\n")
    blocks, i = [], 0
    while i < len(lines):
        ln = lines[i]
        s  = ln.strip()
        if not s:
            i += 1; continue
        if s == "---":
            i += 1; continue                      # separator, replaced by heading rules
        m = re.match(r'^(#{1,3})\s+(.*)$', s)
        if m:
            blocks.append({"t": "h%d" % len(m.group(1)), "text": m.group(2).strip()})
            i += 1; continue
        if s.startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i].strip()); i += 1
            cells = [[c.strip() for c in row.strip("|").split("|")] for row in tbl]
            cells = [r for r in cells if not all(re.fullmatch(r':?-{2,}:?', c) for c in r)]
            blocks.append({"t": "table", "head": cells[0], "rows": cells[1:]})
            continue
        if re.match(r'^\d+\.\s', s):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                items.append(re.sub(r'^\d+\.\s+', '', lines[i].strip())); i += 1
            blocks.append({"t": "ol", "items": items})
            continue
        if s.startswith("- "):
            items = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                indent = len(lines[i]) - len(lines[i].lstrip())
                items.append({"lvl": 1 if indent >= 2 else 0,
                              "text": lines[i].strip()[2:].strip()})
                i += 1
            blocks.append({"t": "ul", "items": items})
            continue
        para = []
        while i < len(lines) and lines[i].strip() and not re.match(
                r'^(#{1,3}\s|\||-\s|\d+\.\s|---$)', lines[i].strip()):
            para.append(lines[i].strip()); i += 1
        blocks.append({"t": "p", "text": " ".join(para)})
    return blocks

# ------------------------------------------------- structure into front + parts
def structure(blocks):
    """Split into front matter and top-level sections; group Tier 1 callouts."""
    front, sections, cur = [], [], None
    for b in blocks:
        if b["t"] == "h1":
            if cur is None and not sections and not front:
                front.append(b); continue          # document title
            cur = {"title": b["text"], "blocks": []}
            sections.append(cur); continue
        (cur["blocks"] if cur else front).append(b)
    return front, sections

def group_callouts(sec):
    """In Tier 1, wrap each list (plus its ':' lead-in) in a callout block."""
    out, blks = [], sec["blocks"]
    i = 0
    while i < len(blks):
        b = blks[i]
        if b["t"] in ("ol", "ul"):
            inner = [b]
            if out and out[-1]["t"] == "p" and out[-1]["text"].rstrip().endswith(":"):
                inner.insert(0, out.pop())
            out.append({"t": "callout", "blocks": inner})
        else:
            out.append(b)
        i += 1
    sec["blocks"] = out
    return sec

def slug(s):
    return "s-" + re.sub(r'[^a-z0-9]+', '-', s.lower()).strip('-')

def build_model():
    md = open(SRC, encoding="utf-8").read()
    front, sections = structure(parse(md))
    for sec in sections:
        if sec["title"].upper().startswith("TIER 1"):
            group_callouts(sec)
        sec["id"]  = slug(sec["title"])
        sec["sub"] = next((v for k, v in TIER_SUB.items()
                           if sec["title"].upper().startswith(k)), None)
        for b in sec["blocks"]:
            if b["t"] == "h2":
                b["id"] = slug(b["text"])
    return front, sections

if __name__ == "__main__":
    f, s = build_model()
    print("front blocks:", len(f))
    for sec in s:
        n = sum(1 for b in sec["blocks"] if b["t"] == "h2")
        c = sum(1 for b in sec["blocks"] if b["t"] == "callout")
        t = sum(1 for b in sec["blocks"] if b["t"] == "table")
        print(f"  {sec['title'][:48]:50s} h2={n:2d} callouts={c:2d} tables={t}")

# ===================================================================== HTML/PDF
NAVY = "#1F3A5F"

CSS = """
@page {
  size: letter;
  margin: 0.7in 1.25in 0.75in 1.0in;   /* wide right margin for pen notes */
  @bottom-left {
    content: "%(exam)s  |  %(title)s";
    font-family: "Liberation Sans", sans-serif; font-size: 8pt; color: #555;
    vertical-align: top; padding-top: 10pt;
  }
  @bottom-right {
    content: "Page " counter(page) " of " counter(pages);
    font-family: "Liberation Sans", sans-serif; font-size: 8pt; color: #555;
    vertical-align: top; padding-top: 10pt;
  }
}

body { font-family: "Liberation Serif", serif; font-size: 11pt; line-height: 1.36;
       color: #000; margin: 0; hyphens: none; }
p { margin: 0 0 5.5pt 0; orphans: 2; widows: 2; }

.doctitle { font-family: "Liberation Sans", sans-serif; font-size: 21pt;
            font-weight: bold; color: %(navy)s; margin: 0 0 4pt 0; line-height: 1.2; }
.docmeta  { font-size: 10pt; color: #333; margin: 0 0 2pt 0;
            border-top: 2pt solid %(navy)s; padding-top: 6pt; }
.rule-b   { border-bottom: 0.75pt solid #B4B4B4; margin: 9pt 0 8pt 0; }

h1.tier { font-family: "Liberation Sans", sans-serif; font-size: 18pt; font-weight: bold;
          color: %(navy)s; margin: 0; padding-bottom: 5pt;
          border-bottom: 2pt solid %(navy)s; break-before: page; break-after: avoid;
          line-height: 1.22; }
.tiersub { font-size: 10.5pt; font-style: italic; color: #333; line-height: 1.4;
           margin: 5pt 0 8pt 0; break-after: avoid; }
.tiergap { margin-bottom: 12pt; }
h1.tier.flow { break-before: auto; margin-top: 15pt; }
.keep { break-inside: avoid; }

h2 { font-family: "Liberation Sans", sans-serif; font-size: 12.5pt; font-weight: bold;
     color: %(navy)s; margin: 8.5pt 0 4pt 0; padding-bottom: 3pt;
     border-bottom: 0.5pt solid #C4C4C4; break-after: avoid; line-height: 1.28; }
h3 { font-family: "Liberation Sans", sans-serif; font-size: 10.8pt; font-weight: bold;
     color: #000; margin: 8pt 0 4pt 0; break-after: avoid; line-height: 1.28; }

ol, ul { margin: 0 0 8pt 0; padding-left: 1.45em; break-inside: avoid; }
li { margin-bottom: 2pt; padding-left: 2pt; }
li:last-child { margin-bottom: 0; }
ul ul { margin: 2pt 0 0 0; }

.callout { background: #F1F1F1; border: 0.75pt solid #9A9A9A; padding: 5pt 9pt 5pt 8pt;
           margin: 4.5pt 0 5pt 0; break-inside: avoid; }
.callout > p:first-child { margin-bottom: 6pt; }
.callout > ol:last-child, .callout > ul:last-child { margin-bottom: 0; }

table { width: 100%%; border-collapse: collapse; margin: 7pt 0 8pt 0;
        font-size: 10pt; break-inside: avoid; }
thead { display: table-header-group; }
th { background: #E6E6E6; font-family: "Liberation Sans", sans-serif; font-size: 9.5pt;
     font-weight: bold; text-align: left; border: 0.75pt solid #7A7A7A;
     padding: 3.5pt 6pt; line-height: 1.22; }
td { border: 0.5pt solid #A8A8A8; padding: 3pt 6pt; vertical-align: top;
     line-height: 1.28; }

/* ---- table of contents: two columns, page numbers flush right ---- */
.toch { font-family: "Liberation Sans", sans-serif; font-size: 13pt; font-weight: bold;
        color: %(navy)s; margin: 11pt 0 6pt 0; padding-bottom: 3pt;
        border-bottom: 0.5pt solid #C4C4C4; }
.toc { column-count: 2; column-gap: 26pt; }
a.toc-entry { display: block; position: relative; text-decoration: none; color: #000;
              font-size: 9.5pt; line-height: 1.32; margin-bottom: 2pt;
              padding-right: 20pt; break-inside: avoid; }
a.toc-entry::after { content: target-counter(attr(href), page);
                     position: absolute; right: 0; top: 0;
                     font-family: "Liberation Sans", sans-serif; font-size: 9pt; }
a.toc-entry.lvl1 { font-family: "Liberation Sans", sans-serif; font-weight: bold;
                   font-size: 9.5pt; color: %(navy)s; margin-top: 8pt;
                   break-after: avoid; }
a.toc-entry.lvl1:first-child { margin-top: 0; }
a.toc-entry.lvl1::after { color: %(navy)s; }
a.toc-entry.lvl2 { padding-left: 13pt; }
""" % {"exam": EXAM, "title": TITLE, "navy": NAVY}


def esc(s):
    return _html.escape(s, quote=False)


def inline_html(s):
    out = []
    for txt, b, i in parse_inline(s):
        t = esc(txt)
        if b: t = "<strong>%s</strong>" % t
        if i: t = "<em>%s</em>" % t
        out.append(t)
    return "".join(out)


ATOMIC = ("table", "callout", "ol", "ul")

def wrap_keeps(blocks):
    """Bind each heading to the atomic block that follows so headings never orphan."""
    out, i = [], 0
    while i < len(blocks):
        b = blocks[i]
        if b["t"] in ("h1sec", "h2", "h3") and i + 1 < len(blocks) \
                and blocks[i + 1]["t"] in ATOMIC:
            out.append({"t": "keep", "blocks": [b, blocks[i + 1]]})
            i += 2
        else:
            out.append(b); i += 1
    return out


def blocks_html(blocks):
    h = []
    for b in blocks:
        t = b["t"]
        if t == "p":
            h.append("<p>%s</p>" % inline_html(b["text"]))
        elif t == "h2":
            h.append('<h2 id="%s">%s</h2>' % (b["id"], inline_html(b["text"])))
        elif t == "h3":
            h.append("<h3>%s</h3>" % inline_html(b["text"]))
        elif t == "ol":
            h.append("<ol>%s</ol>" % "".join(
                "<li>%s</li>" % inline_html(x) for x in b["items"]))
        elif t == "ul":
            h.append(ul_html(b["items"]))
        elif t == "keep":
            h.append('<div class="keep">%s</div>' % blocks_html(b["blocks"]))
        elif t == "h1sec":
            h.append('<h1 class="%s" id="%s">%s</h1>%s'
                     % (b["cls"], b["id"], esc(b["text"]),
                        ('<div class="tiersub">%s</div>' % esc(b["sub"])) if b["sub"]
                        else '<div class="tiergap"></div>'))
        elif t == "callout":
            h.append('<div class="callout">%s</div>' % blocks_html(b["blocks"]))
        elif t == "table":
            head = "".join("<th>%s</th>" % inline_html(c) for c in b["head"])
            rows = "".join("<tr>%s</tr>" % "".join(
                "<td>%s</td>" % inline_html(c) for c in r) for r in b["rows"])
            h.append("<table><thead><tr>%s</tr></thead><tbody>%s</tbody></table>"
                     % (head, rows))
    return "".join(h)


def ul_html(items):
    """Render a flat item list with lvl markers into nested <ul>."""
    h, i = [], 0
    while i < len(items):
        it = items[i]
        if it["lvl"] == 0:
            sub = []
            j = i + 1
            while j < len(items) and items[j]["lvl"] == 1:
                sub.append(items[j]); j += 1
            inner = ("<ul>%s</ul>" % "".join(
                "<li>%s</li>" % inline_html(s["text"]) for s in sub)) if sub else ""
            h.append("<li>%s%s</li>" % (inline_html(it["text"]), inner))
            i = j
        else:
            h.append("<li>%s</li>" % inline_html(it["text"])); i += 1
    return "<ul>%s</ul>" % "".join(h)


def build_html(front, sections):
    p = ['<html><head><meta charset="utf-8"><title>%s</title></head><body>' % esc(TITLE)]
    # --- page 1: title, meta, TOC, how-to-use
    p.append('<div class="doctitle">%s</div>' % esc(TITLE))
    for b in front[1:]:
        if b["t"] == "p" and b["text"].startswith("**Source:**"):
            p.append('<div class="docmeta">%s</div>' % inline_html(b["text"]))
    p.append('<div class="toch">Contents</div><div class="toc">')
    for sec in sections:
        p.append('<a class="toc-entry lvl1" href="#%s">%s</a>' % (sec["id"], esc(sec["title"])))
        for b in sec["blocks"]:
            if b["t"] == "h2":
                p.append('<a class="toc-entry lvl2" href="#%s">%s</a>'
                         % (b["id"], esc(re.sub(r'\*\*(.+?)\*\*', r'\1', b["text"]))))
    p.append("</div>")
    p.append('<div class="rule-b"></div>')
    for b in front[1:]:
        if not (b["t"] == "p" and b["text"].startswith("**Source:**")):
            p.append(blocks_html([b]))
    # --- tier sections
    for sec in sections:
        head = {"t": "h1sec", "cls": "tier" if sec["sub"] else "tier flow",
                "id": sec["id"], "text": sec["title"], "sub": sec["sub"]}
        p.append(blocks_html(wrap_keeps([head] + sec["blocks"])))
    p.append("</body></html>")
    return "".join(p)


def build_pdf():
    from weasyprint import HTML, CSS as WCSS
    front, sections = build_model()
    html = build_html(front, sections)
    open(os.path.join(OUT, "_ucmj.html"), "w", encoding="utf-8").write(html)
    doc = HTML(string=html, base_url=OUT).render(stylesheets=[WCSS(string=CSS)])
    path = os.path.join(OUT, "UCMJ_Study_Guide.pdf")
    doc.write_pdf(path)
    return path, len(doc.pages)

# ======================================================================== DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_BREAK, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY_RGB = RGBColor(0x1F, 0x3A, 0x5F)
SERIF, SANS = "Times New Roman", "Arial"
TEXT_W = 6.25   # inches: 8.5 - 1.0 left - 1.25 right


# OOXML requires child elements in strict schema sequence; append() alone
# produces a file Word and LibreOffice both refuse to open.
_ORDERS = {
 "pPr": ["pStyle","keepNext","keepLines","pageBreakBefore","framePr","widowControl",
         "numPr","suppressLineNumbers","pBdr","shd","tabs","suppressAutoHyphens","kinsoku",
         "wordWrap","overflowPunct","topLinePunct","autoSpaceDE","autoSpaceDN","bidi",
         "adjustRightInd","snapToGrid","spacing","ind","contextualSpacing","mirrorIndents",
         "suppressOverlap","jc","textDirection","textAlignment","textboxTightWrap",
         "outlineLvl","divId","cnfStyle","rPr","sectPr","pPrChange"],
 "tcPr": ["cnfStyle","tcW","gridSpan","hMerge","vMerge","tcBorders","shd","noWrap","tcMar",
          "textDirection","tcFitText","vAlign","hideMark","headers","cellIns","cellDel",
          "cellMerge","tcPrChange"],
 "tblPr": ["tblStyle","tblpPr","tblOverlap","bidiVisual","tblStyleRowBandSize",
           "tblStyleColBandSize","tblW","jc","tblCellSpacing","tblInd","tblBorders","shd",
           "tblLayout","tblCellMar","tblLook","tblCaption","tblDescription","tblPrChange"],
 "trPr": ["cnfStyle","divId","gridBefore","gridAfter","wBefore","wAfter","cantSplit",
          "trHeight","tblHeader","tblCellSpacing","jc","hidden","ins","del","trPrChange"],
 "sectPr": ["headerReference","footerReference","footnotePr","endnotePr","type","pgSz",
            "pgMar","paperSrc","pgBorders","lnNumType","pgNumType","cols","formProt",
            "vAlign","noEndnote","titlePg","textDirection","bidi","rtlGutter","docGrid",
            "printerSettings","sectPrChange"],
 "settings": ["writeProtection","view","zoom","removePersonalInformation",
              "removeDateAndTime","doNotDisplayPageBoundaries","displayBackgroundShape",
              "printPostScriptOverText","printFractionalCharacterWidth","printFormsData",
              "embedTrueTypeFonts","embedSystemFonts","saveSubsetFonts","saveFormsData",
              "mirrorMargins","alignBordersAndEdges","bordersDoNotSurroundHeader",
              "bordersDoNotSurroundFooter","gutterAtTop","hideSpellingErrors",
              "hideGrammaticalErrors","activeWritingStyle","proofState","formsDesign",
              "attachedTemplate","linkStyles","stylePaneFormatFilter","stylePaneSortMethod",
              "documentType","mailMerge","revisionView","trackChanges","doNotTrackMoves",
              "doNotTrackFormatting","documentProtection","autoFormatOverride",
              "styleLockTheme","styleLockQFSet","defaultTabStop","autoHyphenation",
              "consecutiveHyphenLimit","hyphenationZone","doNotHyphenateCaps","showEnvelope",
              "summaryLength","clickAndTypeStyle","defaultTableStyle","evenAndOddHeaders",
              "bookFoldRevPrinting","bookFoldPrinting","bookFoldPrintingSheets",
              "drawingGridHorizontalSpacing","drawingGridVerticalSpacing",
              "displayHorizontalDrawingGridEvery","displayVerticalDrawingGridEvery",
              "doNotUseMarginsForDrawingGridOrigin","drawingGridHorizontalOrigin",
              "drawingGridVerticalOrigin","doNotShadeFormData","noPunctuationKerning",
              "characterSpacingControl","printTwoOnOne","strictFirstAndLastChars",
              "noLineBreaksAfter","noLineBreaksBefore","savePreviewPicture",
              "doNotValidateAgainstSchema","saveInvalidXml","ignoreMixedContent",
              "alwaysShowPlaceholderText","doNotDemarcateInvalidXml","saveXmlDataOnly",
              "useXSLTWhenSaving","saveThroughXslt","showXMLTags",
              "alwaysMergeEmptyNamespace","updateFields","hdrShapeDefaults","footnotePr",
              "endnotePr","compat","docVars","rsids","mathPr","themeFontLang",
              "clrSchemeMapping","doNotIncludeSubdocsInStats","doNotAutoCompressPictures",
              "forceUpgrade","captions","readModeInkLockDown","smartTagType",
              "shapeDefaults","doNotEmbedSmartTags","decimalSymbol","listSeparator"],
}


def _loc(tag):
    return tag.split("}")[-1]


def ins(parent, el):
    """Insert el into parent at its schema-correct position."""
    order = _ORDERS.get(_loc(parent.tag))
    name = _loc(el.tag)
    if not order or name not in order:
        parent.append(el); return el
    idx = order.index(name)
    for child in list(parent):
        cn = _loc(child.tag)
        if cn in order and order.index(cn) > idx:
            child.addprevious(el); return el
    parent.append(el); return el


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def set_font(run, name, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = _el('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), name)
    if size is not None:   run.font.size = Pt(size)
    if bold is not None:   run.bold = bold
    if italic is not None: run.italic = italic
    if color is not None:  run.font.color.rgb = color


def shade(el, hexfill):
    """Apply solid shading to a tc or p element."""
    pr = el.get_or_add_tcPr() if el.tag == qn('w:tc') else el.get_or_add_pPr()
    ins(pr, _el('w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': hexfill}))


def cell_margins(cell, top, start, bottom, end):
    tcpr = cell._tc.get_or_add_tcPr()
    m = _el('w:tcMar')
    for tag, v in (('w:top', top), ('w:start', start), ('w:bottom', bottom), ('w:end', end)):
        m.append(_el(tag, **{'w:w': str(int(v * 1440)), 'w:type': 'dxa'}))
    ins(tcpr, m)


def table_borders(table, sz, color, inside=True):
    tblpr = table._tbl.tblPr
    b = _el('w:tblBorders')
    edges = ['w:top', 'w:left', 'w:bottom', 'w:right']
    if inside: edges += ['w:insideH', 'w:insideV']
    for e in edges:
        b.append(_el(e, **{'w:val': 'single', 'w:sz': str(sz),
                           'w:space': '0', 'w:color': color}))
    ins(tblpr, b)


def add_field(par, instr):
    r = par.add_run(); r._r.append(_el('w:fldChar', **{'w:fldCharType': 'begin'}))
    r = par.add_run(); t = _el('w:instrText', **{'xml:space': 'preserve'})
    t.text = instr; r._r.append(t)
    r = par.add_run(); r._r.append(_el('w:fldChar', **{'w:fldCharType': 'separate'}))
    r = par.add_run("1")
    r = par.add_run(); r._r.append(_el('w:fldChar', **{'w:fldCharType': 'end'}))


def set_cols(section, num, space_in=0.36):
    sectpr = section._sectPr
    cols = sectpr.find(qn('w:cols'))
    if cols is None:
        cols = ins(sectpr, _el('w:cols'))
    cols.set(qn('w:num'), str(num))
    cols.set(qn('w:space'), str(int(space_in * 1440)))
    cols.set(qn('w:equalWidth'), '1')


def page_setup(section):
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin, section.bottom_margin = Inches(0.7), Inches(0.75)
    section.left_margin, section.right_margin = Inches(1.0), Inches(1.25)
    section.footer_distance = Inches(0.4)


def strip_md(s):
    return re.sub(r'\*\*(.+?)\*\*', r'\1', re.sub(r'\*(.+?)\*', r'\1', s))


def pdf_page_map(pdf_path, front, sections):
    """Locate each heading in the rendered PDF so the DOCX TOC caches real numbers."""
    import pymupdf
    doc = pymupdf.open(pdf_path)
    pages = [p.get_text() for p in doc]

    def find(txt, after=0):
        needle = " ".join(strip_md(txt).split())
        for i in range(after, len(pages)):
            hay = " ".join(pages[i].split())
            if needle in hay:
                return i + 1
        return 1

    out, cursor = {}, 0
    for sec in sections:
        n = find(sec["title"], cursor)
        # the TOC on page 1 also contains these strings; skip past it
        if n == 1: n = find(sec["title"], 1)
        out[sec["id"]] = n; cursor = max(0, n - 1)
        for b in sec["blocks"]:
            for bb in (b["blocks"] if b["t"] == "keep" else [b]):
                if bb["t"] == "h2":
                    m = find(bb["text"], cursor)
                    if m == 1: m = find(bb["text"], 1)
                    out[bb["id"]] = m; cursor = max(0, m - 1)
    return out


def runs_into(par, text, base=SERIF, size=11, color=None):
    for txt, b, i in parse_inline(text):
        r = par.add_run(txt)
        set_font(r, base, size, bold=b or None, italic=i or None, color=color)
    if not text:
        set_font(par.add_run(""), base, size)


def add_list_par(container, text, number=None, level=0, size=11):
    par = container.add_paragraph()
    pf = par.paragraph_format
    left = 0.30 + 0.30 * level
    pf.left_indent, pf.first_line_indent = Inches(left), Inches(-0.22)
    pf.space_after, pf.space_before = Pt(1.5), Pt(0)
    pf.line_spacing = 1.36
    pf.tab_stops.add_tab_stop(Inches(left))
    marker = number if number else ("•" if level == 0 else "◦")
    set_font(par.add_run(marker + "\t"), SERIF, size)
    runs_into(par, text, size=size)
    return par


def outline(par, lvl):
    """Mark heading level so a Word TOC field can find it."""
    ins(par._p.get_or_add_pPr(), _el('w:outlineLvl', **{'w:val': str(lvl)}))


def cant_split(row):
    ins(row._tr.get_or_add_trPr(), _el('w:cantSplit'))


def h_border_top(par, color="B4B4B4", sz=6):
    pr = par._p.get_or_add_pPr()
    b = _el('w:pBdr'); b.append(_el('w:top', **{'w:val': 'single', 'w:sz': str(sz),
                                                'w:space': '8', 'w:color': color}))
    ins(pr, b)


def h_border(par, color="C4C4C4", sz=4):
    pr = par._p.get_or_add_pPr()
    b = _el('w:pBdr'); b.append(_el('w:bottom', **{'w:val': 'single', 'w:sz': str(sz),
                                                   'w:space': '2', 'w:color': color}))
    ins(pr, b)


def render_blocks(doc, blocks):
    for b in blocks:
        t = b["t"]
        if t == "keep":
            render_blocks(doc, b["blocks"])
        elif t == "h1sec":
            par = doc.add_paragraph()
            if b["cls"] == "tier":
                par.paragraph_format.page_break_before = True
            pf = par.paragraph_format
            pf.space_before, pf.space_after = Pt(0 if b["cls"] == "tier" else 14), Pt(4)
            pf.keep_with_next = True
            set_font(par.add_run(b["text"]), SANS, 18, bold=True, color=NAVY_RGB)
            outline(par, 0)
            h_border(par, color="1F3A5F", sz=16)
            sub = doc.add_paragraph()
            sub.paragraph_format.space_before = Pt(5)
            sub.paragraph_format.space_after = Pt(8)
            sub.paragraph_format.keep_with_next = True
            if b["sub"]:
                set_font(sub.add_run(b["sub"]), SERIF, 10.5, italic=True,
                         color=RGBColor(0x33, 0x33, 0x33))
        elif t == "h2":
            par = doc.add_paragraph()
            pf = par.paragraph_format
            pf.space_before, pf.space_after = Pt(7.5), Pt(3.5)
            pf.keep_with_next = True
            set_font(par.add_run(strip_md(b["text"])), SANS, 12.5, bold=True, color=NAVY_RGB)
            outline(par, 1)
            h_border(par)
        elif t == "h3":
            par = doc.add_paragraph()
            pf = par.paragraph_format
            pf.space_before, pf.space_after = Pt(8), Pt(4)
            pf.keep_with_next = True
            set_font(par.add_run(strip_md(b["text"])), SANS, 10.8, bold=True)
        elif t == "p":
            par = doc.add_paragraph()
            par.paragraph_format.space_after = Pt(5)
            par.paragraph_format.line_spacing = 1.36
            runs_into(par, b["text"])
        elif t == "ol":
            for i, it in enumerate(b["items"], 1):
                add_list_par(doc, it, number="%d." % i)
        elif t == "ul":
            for it in b["items"]:
                add_list_par(doc, it["text"], level=it["lvl"])
        elif t == "callout":
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.autofit = False
            cell = tbl.cell(0, 0)
            cell.width = Inches(TEXT_W)
            shade(cell._tc, "F1F1F1")
            cell_margins(cell, 0.07, 0.11, 0.07, 0.11)
            table_borders(tbl, 6, "9A9A9A", inside=False)
            cant_split(tbl.rows[0])
            cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
            render_blocks(cell, b["blocks"])
            for p_ in cell.paragraphs:
                p_.paragraph_format.space_after = Pt(2)
            cell.paragraphs[-1].paragraph_format.space_after = Pt(0)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = 1
            set_font(spacer.add_run(""), SERIF, 3)
        elif t == "table":
            tbl = doc.add_table(rows=1, cols=len(b["head"]))
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            table_borders(tbl, 4, "A8A8A8")
            hdr = tbl.rows[0]
            trpr = hdr._tr.get_or_add_trPr(); ins(trpr, _el('w:tblHeader'))
            for j, c in enumerate(b["head"]):
                cell = hdr.cells[j]
                shade(cell._tc, "E6E6E6")
                cell_margins(cell, 0.035, 0.06, 0.035, 0.06)
                par = cell.paragraphs[0]
                par.paragraph_format.space_after = Pt(0)
                par.paragraph_format.line_spacing = 1.15
                set_font(par.add_run(strip_md(c)), SANS, 9.5, bold=True)
            for row in b["rows"]:
                cells = tbl.add_row().cells
                for j, c in enumerate(row):
                    cell_margins(cells[j], 0.03, 0.06, 0.03, 0.06)
                    par = cells[j].paragraphs[0]
                    par.paragraph_format.space_after = Pt(0)
                    par.paragraph_format.line_spacing = 1.2
                    runs_into(par, c, size=10)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = 1
            set_font(spacer.add_run(""), SERIF, 4)


def build_docx(pdf_path=None):
    front, sections = build_model()
    pmap = pdf_page_map(pdf_path or os.path.join(OUT, "UCMJ_Study_Guide.pdf"),
                        front, sections) if pdf_path is not False else {}

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name, st.font.size = SERIF, Pt(11)
    st.element.rPr.rFonts.set(qn('w:ascii'), SERIF)
    st.element.rPr.rFonts.set(qn('w:hAnsi'), SERIF)
    st.paragraph_format.line_spacing = 1.36
    st.paragraph_format.space_after = Pt(5.5)

    s0 = doc.sections[0]
    page_setup(s0); set_cols(s0, 1)

    # running footer: exam number left, page N of M right
    fp = s0.footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(TEXT_W), WD_TAB_ALIGNMENT.RIGHT)
    grey = RGBColor(0x55, 0x55, 0x55)
    set_font(fp.add_run("%s  |  %s\t" % (EXAM, TITLE)), SANS, 8, color=grey)
    set_font(fp.add_run("Page "), SANS, 8, color=grey)
    add_field(fp, " PAGE ")
    set_font(fp.add_run(" of "), SANS, 8, color=grey)
    add_field(fp, " NUMPAGES ")
    for r in fp.runs:
        set_font(r, SANS, 8, color=grey)

    # ---- page 1: title + meta
    t = doc.add_paragraph(); t.paragraph_format.space_after = Pt(3)
    set_font(t.add_run(TITLE), SANS, 21, bold=True, color=NAVY_RGB)
    h_border(t, color="1F3A5F", sz=18)
    for b in front[1:]:
        if b["t"] == "p" and b["text"].startswith("**Source:**"):
            m = doc.add_paragraph()
            m.paragraph_format.space_before = Pt(4)
            m.paragraph_format.space_after = Pt(2)
            runs_into(m, b["text"], size=10)
    c = doc.add_paragraph()
    c.paragraph_format.space_before, c.paragraph_format.space_after = Pt(8), Pt(4)
    set_font(c.add_run("Contents"), SANS, 13, bold=True, color=NAVY_RGB)
    h_border(c)

    # ---- TOC: single column, dot leaders, live field over cached entries
    entries = []
    for sec in sections:
        entries.append((sec["title"], sec["id"], 1))
        for b in sec["blocks"]:
            for bb in (b["blocks"] if b["t"] == "keep" else [b]):
                if bb["t"] == "h2":
                    entries.append((strip_md(bb["text"]), bb["id"], 2))

    toc_pars = []
    for text, tid, lvl in entries:
        par = doc.add_paragraph()
        pf = par.paragraph_format
        pf.space_before = Pt(2 if lvl == 1 else 0)
        pf.space_after, pf.line_spacing = Pt(0.5), 1.0
        pf.left_indent = Inches(0 if lvl == 1 else 0.20)
        pf.first_line_indent = Inches(0)
        pf.tab_stops.add_tab_stop(Inches(TEXT_W), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        pf.keep_with_next = (lvl == 1)
        toc_pars.append((par, text, tid, lvl))

    for idx, (par, text, tid, lvl) in enumerate(toc_pars):
        if idx == 0:
            r = par.add_run(); r._r.append(_el('w:fldChar', **{'w:fldCharType': 'begin'}))
            r = par.add_run(); it = _el('w:instrText', **{'xml:space': 'preserve'})
            it.text = ' TOC \\o "1-2" \\h \\z \\u '; r._r.append(it)
            r = par.add_run(); r._r.append(_el('w:fldChar', **{'w:fldCharType': 'separate'}))
        if lvl == 1:
            set_font(par.add_run(text), SANS, 9, bold=True, color=NAVY_RGB)
            set_font(par.add_run("\t" + str(pmap.get(tid, 1))), SANS, 8.5, bold=True,
                     color=NAVY_RGB)
        else:
            set_font(par.add_run(text), SERIF, 9)
            set_font(par.add_run("\t" + str(pmap.get(tid, 1))), SANS, 8.5)
        if idx == len(toc_pars) - 1:
            r = par.add_run(); r._r.append(_el('w:fldChar', **{'w:fldCharType': 'end'}))

    # separator rule rides on the next paragraph rather than an empty one
    idx = len(doc.paragraphs)
    for b in front[1:]:
        if not (b["t"] == "p" and b["text"].startswith("**Source:**")):
            render_blocks(doc, [b])
    intro = doc.paragraphs[idx]
    intro.paragraph_format.space_before = Pt(8)
    h_border_top(intro)

    for sec in sections:
        head = {"t": "h1sec", "cls": "tier" if sec["sub"] else "tier flow",
                "id": sec["id"], "text": sec["title"], "sub": sec["sub"]}
        render_blocks(doc, wrap_keeps([head] + sec["blocks"]))

    # ask Word to refresh the TOC page numbers on open
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is None:
        ins(settings, _el('w:updateFields', **{'w:val': 'true'}))

    path = os.path.join(OUT, "UCMJ_Study_Guide.docx")
    doc.save(path)
    return path
